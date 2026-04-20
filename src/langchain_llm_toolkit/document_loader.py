import os
import magic
from typing import List
from langchain_core.documents import Document
from langchain_llm_toolkit.markdown_loader import MarkdownLoader


class DocumentLoader:
    """文档加载器

    支持的格式：
    - PDF
    - TXT
    - Markdown（增强版）
    - DOCX
    """

    def __init__(
        self,
        markdown_split_by_heading: bool = True,
        markdown_extract_metadata: bool = True,
        markdown_extract_code_blocks: bool = True,
        markdown_extract_tables: bool = True,
        markdown_min_heading_level: int = 1,
        markdown_max_heading_level: int = 6,
    ):
        """
        初始化文档加载器

        Args:
            markdown_split_by_heading: 是否按标题分割 Markdown 文档
            markdown_extract_metadata: 是否提取 Markdown 元数据
            markdown_extract_code_blocks: 是否提取 Markdown 代码块
            markdown_extract_tables: 是否提取 Markdown 表格
            markdown_min_heading_level: Markdown 最小标题级别
            markdown_max_heading_level: Markdown 最大标题级别
        """
        self.mime = magic.Magic(mime=True)
        self.markdown_loader = MarkdownLoader(
            split_by_heading=markdown_split_by_heading,
            extract_metadata=markdown_extract_metadata,
            extract_code_blocks=markdown_extract_code_blocks,
            extract_tables=markdown_extract_tables,
            min_heading_level=markdown_min_heading_level,
            max_heading_level=markdown_max_heading_level,
        )

    def load_document(self, file_path: str) -> List[Document]:
        """加载文档并返回文档对象列表"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        file_extension = os.path.splitext(file_path)[1].lower()

        if file_extension == ".pdf":
            return self._load_pdf(file_path)
        elif file_extension == ".txt":
            return self._load_txt(file_path)
        elif file_extension == ".md":
            return self._load_markdown(file_path)
        elif file_extension == ".docx":
            return self._load_docx(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {file_extension}")

    def load_documents(self, file_paths: List[str]) -> List[Document]:
        """批量加载多个文档"""
        all_documents = []
        for file_path in file_paths:
            try:
                documents = self.load_document(file_path)
                all_documents.extend(documents)
            except Exception as e:
                print(f"加载文件 {file_path} 失败: {e}")
        return all_documents

    def _load_pdf(self, file_path: str) -> List[Document]:
        """加载 PDF 文档"""
        try:
            import pypdf

            documents = []
            with open(file_path, "rb") as file:
                reader = pypdf.PdfReader(file)
                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text:
                        doc = Document(
                            page_content=text,
                            metadata={
                                "source": file_path,
                                "page": page_num + 1,
                                "total_pages": len(reader.pages),
                                "type": "pdf",
                            },
                        )
                        documents.append(doc)
            return documents
        except ImportError:
            raise ImportError("需要安装 pypdf 来处理 PDF 文件")

    def _load_txt(self, file_path: str) -> List[Document]:
        """加载 TXT 文档"""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
            text = file.read()
            doc = Document(page_content=text, metadata={"source": file_path, "type": "txt"})
            return [doc]

    def _load_markdown(self, file_path: str) -> List[Document]:
        """加载 Markdown 文档（增强版）"""
        return self.markdown_loader.load(file_path)

    def _load_docx(self, file_path: str) -> List[Document]:
        """加载 DOCX 文档"""
        try:
            from docx import Document as DocxDocument

            docx = DocxDocument(file_path)
            text = "\n".join([para.text for para in docx.paragraphs])
            doc = Document(page_content=text, metadata={"source": file_path, "type": "docx"})
            return [doc]
        except ImportError:
            raise ImportError("需要安装 python-docx 来处理 DOCX 文件")

    def get_markdown_outline(self, file_path: str) -> List[dict]:
        """获取 Markdown 文档大纲"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        return self.markdown_loader.get_document_outline(content)

    def get_markdown_statistics(self, file_path: str) -> dict:
        """获取 Markdown 文档统计信息"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        return self.markdown_loader.get_statistics(content)
